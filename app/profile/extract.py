"""Extract plain text from an uploaded resume (PDF, DOCX, Markdown, TXT)."""
import io

SUPPORTED = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/markdown": "text",
    "text/plain": "text",
}
EXT_TO_KIND = {".pdf": "pdf", ".docx": "docx", ".md": "text", ".txt": "text"}


class UnsupportedResume(ValueError):
    pass


def detect_kind(filename: str, content_type: str | None) -> str:
    lower = (filename or "").lower()
    for ext, kind in EXT_TO_KIND.items():
        if lower.endswith(ext):
            return kind
    if content_type in SUPPORTED:
        return SUPPORTED[content_type]
    raise UnsupportedResume(f"unsupported resume type: {filename!r} ({content_type})")


def extract_text(data: bytes, filename: str, content_type: str | None = None) -> str:
    kind = detect_kind(filename, content_type)
    if kind == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    elif kind == "docx":
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
    else:
        text = data.decode("utf-8", errors="replace")

    text = "\n".join(line.rstrip() for line in text.splitlines())
    if not text.strip():
        raise UnsupportedResume("no extractable text found (scanned image PDF?)")
    return text.strip()
